#!/usr/bin/env python3
"""Generate the PeakBuddy Audit & Wearables Integration Plan as a .docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette ----
INK = RGBColor(0x1A, 0x1A, 0x1A)
ACCENT = RGBColor(0x0F, 0x4C, 0x81)      # deep blue
MUTED = RGBColor(0x6B, 0x6B, 0x6B)
CRIT = RGBColor(0xC0, 0x1A, 0x1A)
HIGH = RGBColor(0xC2, 0x5E, 0x00)
MED = RGBColor(0xB8, 0x8A, 0x00)
LOW = RGBColor(0x4A, 0x6B, 0x2A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# base style
base = doc.styles["Normal"]
base.font.name = "Calibri"
base.font.size = Pt(10.5)
base.font.color.rgb = INK

def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def heading(text, level=1, color=ACCENT):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    sizes = {1: 16, 2: 13, 3: 11.5}
    run.font.size = Pt(sizes.get(level, 11))
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "0F4C81")
        pbdr.append(bottom)
        pPr.append(pbdr)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
    return p

def para(text, italic=False, color=None, size=10.5, after=6, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if color:
        r.font.color.rgb = color
    r.font.size = Pt(size)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def table(headers, rows, widths=None, sev_col=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        _shade(hdr[i], "0F4C81")
        para_cell = hdr[i].paragraphs[0]
        run = para_cell.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9.5)
    sev_colors = {"Critical": "F6D4D4", "High": "F8E2CC", "Medium": "F7F0CC", "Low": "E4ECD9"}
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cp = cells[i].paragraphs[0]
            run = cp.add_run(str(val))
            run.font.size = Pt(9)
            if sev_col is not None and i == sev_col and val in sev_colors:
                _shade(cells[i], sev_colors[val])
                run.bold = True
    if widths:
        for i, w in enumerate(widths):
            for c in t.columns[i].cells:
                c.width = Inches(w)
    return t

# =====================================================================
# COVER
# =====================================================================
for _ in range(3):
    doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("PeakBuddy")
r.bold = True
r.font.size = Pt(34)
r.font.color.rgb = ACCENT

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Application Audit & Wearables Integration Plan")
r.font.size = Pt(16)
r.font.color.rgb = INK

s2 = doc.add_paragraph()
s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s2.add_run("Oura • Polar • Garmin")
r.font.size = Pt(12)
r.font.color.rgb = MUTED

for _ in range(6):
    doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line, c in [("Prepared by: Muhammad Asad Khan", INK),
                ("Date: 20 June 2026", MUTED),
                ("Status: Technical assessment — for client review", MUTED),
                ("Scope: TanStack Start (React 19) · Cloudflare Workers · Supabase", MUTED)]:
    rr = meta.add_run(line + "\n")
    rr.font.size = Pt(10.5)
    rr.font.color.rgb = c
doc.add_page_break()

# =====================================================================
# EXECUTIVE SUMMARY
# =====================================================================
heading("1. Executive Summary", 1)
para("PeakBuddy (“Buddy Tracker by Peak Movement”) is a symptom-tracking PWA for "
     "musculoskeletal / movement practitioners and their clients. It is built on TanStack Start "
     "(React 19), deployed to Cloudflare Workers, backed by Supabase (Postgres + Auth + RLS), and "
     "was originally scaffolded with Lovable. It supports three personas — client, practitioner, "
     "and super-admin — with daily check-ins, progress charts, an AI symptom-triage assistant "
     "(“Yves”), program suggestions, and a practitioner alert pipeline.")
para("This document has two parts. Part A is a technical audit of the current application across "
     "architecture, data model, security, and code quality. Part B is a concrete plan to integrate "
     "wearable devices (Oura, Polar, and Garmin) so that objective biometric data — sleep, "
     "heart-rate, HRV, readiness, activity — augments the existing self-reported check-ins.", after=8)

heading("Headline findings", 3)
bullet("the app is well-structured and security-aware in its hardest area (the public AI triage "
       "endpoint), but several privileged server functions bypass row-level security with no "
       "authentication — including one account-takeover vector that should be fixed before any "
       "new feature work.", "Security: ")
bullet("clean, well-typed code with PHI-aware logging, but testing is near-absent and an entire "
       "data-fetching layer (TanStack Query) is installed yet unused.", "Code quality: ")
bullet("no biometric/time-series storage exists today; a new table and provider-sync layer are "
       "required. The codebase has clear, reusable patterns to plug this into.", "Wearables: ")
bullet("Oura and Polar are self-serve OAuth2 and can ship first. Garmin requires partner-program "
       "approval (not self-serve) and its developer program is reportedly on hold in 2026 — it "
       "should be sequenced last or via an aggregator.", "Garmin caveat: ")

# =====================================================================
# PART A
# =====================================================================
doc.add_page_break()
heading("PART A — APPLICATION AUDIT", 1)

heading("2. Architecture & Stack", 2)
table(["Layer", "Technology", "Notes"],
      [["Frontend", "React 19, TanStack Start + Router", "Flat file-based routing (persona.app.feature.tsx)"],
       ["Styling", "Tailwind v4, Radix UI", "Screens use large inline style objects; Tailwind mostly in src/components/ui"],
       ["Data layer", "Supabase JS (anon + service-role)", "Manual useState/useEffect fetch; TanStack Query installed but unused"],
       ["Server logic", "createServerFn (*.functions.ts)", "Zod validation + lazy server-only admin client import"],
       ["Backend", "Supabase Postgres, Auth, RLS", "8 tables, 31 migrations"],
       ["Hosting", "Cloudflare Workers (wrangler)", "src/server.ts entry; nodejs_compat"],
       ["AI", "Anthropic (Claude), Google Gemini via Lovable gateway", "Consent-gated; disclosed in privacy policy"],
       ["Email", "Resend", "Practitioner notifications"]],
      widths=[1.3, 2.3, 3.2])
para("")
bullet("Three persona route groups: client (checkin, timeline, progress, yves, profile), "
       "practitioner (dashboard, clients, program-queue, alerts, settings), and admin.", "Routing: ")
bullet("Two Supabase clients — an anon/RLS browser client and a server-only service-role admin "
       "client that bypasses RLS. The service key is correctly kept out of the browser bundle "
       "(src/integrations/supabase/client.server.ts).", "Clients: ")
bullet("Client-side sessions are unauthenticated localStorage (buddy.client_id); browser reads rely "
       "entirely on Supabase RLS for protection.", "Client identity: ")

heading("3. Data Model", 2)
para("Eight tables, all with row-level security enabled. Summary:")
table(["Table", "Purpose"],
      [["profiles", "Identity + role (super_admin / practitioner / client), FK to auth.users"],
       ["practices", "One row per practitioner: config, POPIA/consent flags, webhooks, Yves toggle"],
       ["clients", "Patient record owned by a practitioner; login_code, program state, yves_ai_consent"],
       ["check_ins", "Self-reported time-series: pain/sleep/stress/energy (integer scales), mood, notes"],
       ["alerts", "Red-flag/triage alerts for practitioners; urgency, webhook_fired, assessment"],
       ["symptom_queries", "Yves AI triage queries + results (urgency, red flags, differential jsonb)"],
       ["programs", "Global catalog of recommendable programs; admin-approval columns"],
       ["platform_settings", "Singleton global config / feature flags (super-admin only)"]],
      widths=[1.5, 5.3])
para("")
heading("Data-model concerns", 3)
bullet("Clients are linked to auth users by lowercased email match in the handle_new_user / "
       "email-confirmed triggers. If the emails differ, the link silently never happens, leaving "
       "auth_user_id NULL and breaking all client-scoped RLS. This failure already occurred and "
       "needed a remediation migration.", "Fragile identity linkage: ")
bullet("current_client_id and is_super_admin exist in both public and private schemas; the public "
       "copies have had EXECUTE repeatedly revoked. An orphaned policy referencing the revoked copy "
       "broke client login once already. Risk of future code calling the wrong copy.", "Duplicate functions / churn: ")
bullet("Several versioned migrations hard-code production UUIDs/emails (delete-by-email, "
       "promote-specific-user, merge duplicates). These are environment-specific and behave oddly on "
       "a fresh database.", "Embedded prod data migrations: ")
bullet("prevent_role_escalation has been rewritten ~4 times; its service-role detection via JWT "
       "claims is brittle across Supabase/PostgREST versions.", "Role-escalation guard brittleness: ")

heading("4. Security Findings", 2)
para("Findings are ranked by severity. The two Critical items were verified directly in source. The "
     "root cause is shared: several server functions call the RLS-bypassing service-role admin client "
     "with no authentication middleware and trust caller-supplied target IDs.", after=8)
table(["ID", "Severity", "Finding", "Location"],
      [["C1", "Critical", "setYvesAiConsent has no auth; flips AI consent for any client by caller-supplied ID (IDOR write + consent/compliance bypass)", "yves-consent.functions.ts:3-26"],
       ["C2", "Critical", "createClientAccount has no auth; trusts practitionerId, and RESETS the password of any existing account by email — account takeover", "clients.functions.ts:16-97 (line 55)"],
       ["H1", "High", "notifyAssignedPractitioner unauthenticated — send attacker text to practitioner email + inject false alerts via any clientId", "notify-practitioner.functions.ts:75-167"],
       ["H2", "High", "getClientYvesAccess unauthenticated RLS-bypassing read; leaks practitioner UUID + flags", "yves-access.functions.ts:3-42"],
       ["H3", "High", "suggestProgram unauthenticated; writes program suggestions for arbitrary clients (AI call itself is consent-gated)", "programs.functions.ts:137-210"],
       ["M1", "Medium", "Triage rate-limit is in-memory — ineffective across Cloudflare isolates; false sense of cost protection", "triage-query.ts:22-37"],
       ["M2", "Medium", "CORS reflects first allowed origin as fallback for disallowed origins (not exploitable; correctness)", "triage-query.ts:12-19"],
       ["M3", "Medium", "Patient text interpolated into LLM prompt — triage output is attacker-influenceable (advisory only)", "triage-query.ts:482"],
       ["L1", "Low", ".env is committed (anon key only — not a leak); stop tracking to prevent future secret commits", ".env"],
       ["L2", "Low", "Two parallel admin-authorization mechanisms (RPC vs profiles.role read) invite drift", "admin-*.functions.ts"],
       ["L3", "Low", "notify-practitioner uses a separate SEED_SERVICE_ROLE_KEY, broadening secret surface", "notify-practitioner.functions.ts:78"]],
      widths=[0.4, 0.8, 3.9, 1.7], sev_col=1)
para("")
heading("Priority remediation (do before wearables work)", 3)
bullet("Add requireSupabaseAuth middleware + caller-ownership verification to the five unauthenticated "
       "admin-client functions (C1, C2, H1, H2, H3). The triage endpoint already demonstrates the "
       "correct pattern (bearer auth via getUser, ownership check, consent check).")
bullet("In createClientAccount: force practitionerId = authenticated caller's id, and never reset an "
       "existing user's password on the “already exists” branch (C2).")
bullet("Move the triage rate-limit to a durable store (Cloudflare KV / Durable Object / DB) keyed on "
       "user id + IP (M1).")

heading("Things already done well", 3)
bullet("Public triage endpoint: proper bearer auth, ownership + consent checks, fail-closed lookups, "
       "length caps, server-built context. Best-hardened code in the repo.")
bullet("Service-role key is server-only (.server.ts), never bundled to the client.")
bullet("PHI-aware central logger redacts known PHI keys and silences debug/info in prod.")
bullet("account-delete is correctly gated and acts only on the caller's own identity.")
bullet("Webhook delivery is https-only with SSRF guards (blocks localhost/RFC1918/link-local).")

heading("5. Code Quality & Testing", 2)
table(["Area", "Assessment"],
      [["Type safety", "Good — all any/as-any confined to generated route tree; no @ts-ignore in hand code"],
       ["Markers", "Clean — no TODO/FIXME/HACK in src/"],
       ["Logging", "Strong — PHI-aware redaction, enforced by ESLint no-console"],
       ["Error handling", "Solid — root error/404 boundaries, out-of-band error capture, generic user copy"],
       ["Offline", "Robust bespoke localStorage queue; never drops check-ins; re-queues on failure"],
       ["Testing", "Weak — only yves.test.ts; no E2E, no server-fn/component tests"],
       ["Data fetching", "TanStack Query installed but unused; manual useEffect fetch duplicated ~29 files"],
       ["Styling", "Heavy inline style objects despite Tailwind being configured"],
       ["CI", "typecheck + lint + vitest on PR; no build step, no E2E, no security scan"]],
      widths=[1.5, 5.3])
para("")
bullet("No vite build in CI — a broken production build can pass. Add a build step.")
bullet("No dependency/security scanning (npm audit / Dependabot / CodeQL / secret scan).")
bullet("Business logic (alert/pattern/webhook orchestration) lives inside the check-in component, "
       "duplicated with offline-queue.ts — hard to test; extract to a service module.")

# =====================================================================
# PART B
# =====================================================================
doc.add_page_break()
heading("PART B — WEARABLES INTEGRATION PLAN", 1)

heading("6. Goals & Use Cases", 2)
para("Augment subjective daily check-ins with objective biometric signals so practitioners get a "
     "fuller, lower-friction picture of client recovery and load:")
bullet("Sleep duration & quality, readiness/recovery scores → contextualize reported sleep_quality and energy.")
bullet("Resting heart-rate and HRV trends → early signal of overtraining, illness, or stress.")
bullet("Activity / steps / workouts → adherence and load management against programs.")
bullet("Biometric-driven alerts → feed the existing practitioner alert pipeline (e.g. HRV crash, poor sleep streak).")

heading("7. Provider Comparison", 2)
table(["", "Oura", "Polar", "Garmin"],
      [["Auth", "OAuth2", "OAuth2 (scope accesslink.read_all)", "OAuth 1.0a (PKCE for newer)"],
       ["Access", "Self-serve", "Self-serve", "Partner program — apply as legal entity"],
       ["Delivery", "Webhooks + REST (v2)", "Webhooks (push) + transaction REST", "Push (callback URLs)"],
       ["Data", "Sleep, readiness, activity, HR, HRV, workouts", "Exercise, daily activity, sleep, HR, HRV", "Activities + broad health metrics"],
       ["History", "Available via REST", "Only last 90 days at registration", "Per partner terms"],
       ["2026 status", "Stable (PATs deprecated Dec 2025)", "Stable", "Connect Dev Program reportedly ON HOLD"],
       ["Effort", "Low", "Medium (transaction model)", "High (approval + OAuth1.0a + uncertainty)"]],
      widths=[1.0, 1.9, 1.95, 1.95])
para("")
para("Recommendation: ship Oura first (lowest effort), Polar second, and treat Garmin as a separate "
     "track — apply for partner access early, and keep a third-party aggregator (e.g. Terra, Spike, "
     "Thryve) as a fallback that covers all three behind one API if direct Garmin access stalls.",
     bold=False, italic=True)

heading("8. Recommended Architecture", 2)
bullet("OAuth handled entirely server-side via new *.functions.ts server functions (authorize + "
       "token exchange + refresh), mirroring the existing clients.functions.ts / yves-access pattern "
       "so provider client-secrets never reach the browser.", "OAuth: ")
bullet("Each provider exposes webhooks/push. Add a Cloudflare-hosted callback route (alongside "
       "src/routes/api/public/triage-query.ts) that verifies the provider signature (HMAC), then "
       "pulls the changed data and upserts it. Poll as a fallback where webhooks are thin.", "Ingestion: ")
bullet("Store provider tokens in a new wearable_connections table with RLS; refresh on a schedule "
       "(Cloudflare Cron / Supabase scheduled function).", "Tokens: ")
bullet("Normalize each provider's payload into a common metric vocabulary before storage so the UI "
       "is provider-agnostic.", "Normalization: ")
bullet("Reuse the existing alert pipeline (insert_alert + fireAlertWebhook) with new alert_type / "
       "pattern values for biometric thresholds — no change to the alerts UI.", "Alerts: ")

heading("9. Data Model Changes", 2)
para("No biometric storage exists today (check_ins holds only manual integer scales). Proposed new tables:")
code("wearable_connections (\n"
     "  id, client_id → clients, provider ('oura'|'polar'|'garmin'),\n"
     "  access_token, refresh_token, token_expires_at, provider_user_id,\n"
     "  scopes, status, connected_at, last_sync_at )")
code("wearable_readings (\n"
     "  id, client_id → clients, provider, metric_type ('sleep'|'hrv'|'rhr'|'steps'|...),\n"
     "  value, unit, recorded_at (device time), ingested_at,\n"
     "  source_id (dedup key), raw jsonb )\n"
     "  index on (client_id, metric_type, recorded_at)")
bullet("RLS: client can read own rows; practitioner/admin read their clients' rows — mirror the "
       "check_ins policies. Writes only via service-role ingestion.")
bullet("Consent: gate connection + sync behind a per-client consent flag, mirroring the existing "
       "yves_ai_consent pattern, and disclose providers in the privacy policy (POPIA / App Store).")
bullet("source_id gives idempotent upserts so re-synced/webhook-replayed samples don’t duplicate.")

heading("10. UI Surfaces", 2)
para("In priority order, reusing existing chart infrastructure:")
bullet("practitioner client-detail — richest surface (metricMeta config, segmented metric "
       "selector, last30 series, Line/Bar/Ring charts). Extend these to add HR/HRV/sleep/steps "
       "series. This is the established extension point.", "1. ")
bullet("client progress — mirror MetricRing/Recharts for a recovery/HRV ring or trend.", "2. ")
bullet("client timeline — add biometric rows to the expandable per-day detail.", "3. ")
bullet("admin client-detail — read-only biometric columns.", "4. ")
bullet("New “Connect a device” screen in client profile/settings to start the OAuth flow and show "
       "connection status.", "5. ")

heading("11. Phased Rollout & Effort", 2)
table(["Phase", "Scope", "Est."],
      [["0. Security fixes", "Remediate C1/C2/H1–H3 before new feature work (prerequisite)", "3–5 d"],
       ["1. Foundation", "wearable_connections + wearable_readings tables, RLS, consent flag, token storage", "3–4 d"],
       ["2. Oura", "OAuth flow, webhook callback, normalization, ingestion, connect UI", "5–7 d"],
       ["3. Display", "Extend practitioner client-detail + client progress charts; biometric alerts", "5–7 d"],
       ["4. Polar", "OAuth + transaction-model ingestion + webhooks", "5–7 d"],
       ["5. Garmin", "Partner application (lead time) + OAuth1.0a + push ingestion, OR aggregator fallback", "7–10 d + approval lead time"],
       ["6. Hardening", "Token refresh cron, retries/backfill, dedup, tests, build+scan in CI", "4–6 d"]],
      widths=[1.4, 4.6, 1.0])
para("")
para("Estimates assume one developer and exclude Garmin partner-approval calendar time (typically a "
     "few business days to apply, but the program’s 2026 hold may extend this). Oura + Polar + "
     "display can realistically ship in ~4–5 weeks after the security prerequisite.", italic=True)

# =====================================================================
# PART C
# =====================================================================
doc.add_page_break()
heading("PART C — ACCESS & CREDENTIALS REQUIRED", 1)
para("To audit, run, deploy, and build the wearables feature, the following access is needed from the "
     "client (GitHub access is already granted).")
heading("Platform access", 3)
table(["System", "What's needed", "Why"],
      [["GitHub", "Granted ✓", "Source code"],
       ["Supabase", "Project member invite (Developer/Admin) + SERVICE_ROLE_KEY", "DB, migrations, RLS, auth, ingestion"],
       ["Lovable", "Workspace invite + LOVABLE_API_KEY", "Build platform + AI gateway (Gemini)"],
       ["Cloudflare", "Account access (or confirm deploys go via Lovable)", "Hosting, Workers, KV/Cron, webhook routes"],
       ["Anthropic", "ANTHROPIC_API_KEY", "Yves triage"],
       ["Resend", "RESEND_API_KEY", "Practitioner email"]],
      widths=[1.3, 3.0, 2.5])
para("")
heading("Wearable developer accounts (client must register as the data controller)", 3)
table(["Provider", "What to set up"],
      [["Oura", "Oura Cloud developer app → client_id/secret, redirect URI, webhook URL"],
       ["Polar", "Polar AccessLink / admin account → client_id/secret, callback URL"],
       ["Garmin", "Apply to Garmin Health/Connect partner program as a legal entity (lead time); OR aggregator account"]],
      widths=[1.3, 5.5])
para("")
heading("Recommended request to the client", 3)
para("“Please (1) invite me to Supabase and share the service-role key; (2) add me to the Lovable "
     "workspace and share the Lovable API key; (3) confirm how deploys work / give Cloudflare access; "
     "(4) share the Anthropic and Resend keys; and (5) for wearables, register developer apps with "
     "Oura and Polar under the company account and start the Garmin partner application. Also: is "
     "there a staging environment separate from production data?”", italic=True)

# =====================================================================
# APPENDIX
# =====================================================================
doc.add_page_break()
heading("Appendix — Key File References", 1)
refs = [
    ("Security", "src/lib/yves-consent.functions.ts, clients.functions.ts, notify-practitioner.functions.ts, yves-access.functions.ts, programs.functions.ts, routes/api/public/triage-query.ts"),
    ("Supabase clients", "src/integrations/supabase/{client.ts, client.server.ts, auth-middleware.ts, auth-attacher.ts}"),
    ("Data model", "supabase/migrations/ (31 files), src/integrations/supabase/types.ts, src/lib/types.ts"),
    ("Check-in / offline", "src/routes/client.app.checkin.tsx, src/lib/offline-queue.ts"),
    ("Display surfaces", "src/routes/practitioner.app.client-detail.$clientId.tsx, client.app.progress.tsx, client.app.timeline.tsx"),
    ("Infra / CI", "wrangler.jsonc, src/server.ts, .github/workflows/ci.yml, src/lib/log.ts, webhooks.ts"),
]
for k, v in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(k + ": ")
    r.bold = True
    r.font.color.rgb = ACCENT
    r.font.size = Pt(10)
    rr = p.add_run(v)
    rr.font.name = "Consolas"
    rr.font.size = Pt(8.5)

para("")
para("Disclaimer: This assessment is based on a static read of the repository at the stated date. "
     "Severity ratings reflect code review; exploitability should be confirmed against the live "
     "deployment with appropriate authorization before remediation sign-off.", italic=True, color=MUTED, size=9)

out = "/Users/asadkhan/work/peakbuddy/PeakBuddy_Audit_and_Wearables_Plan.docx"
doc.save(out)
print("Saved:", out)
